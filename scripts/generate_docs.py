#!/usr/bin/env python3
"""
scripts/generate_docs.py
Gera a documentação completa do projeto ZIYOU Analytics em formato DOCX.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

OUT = Path(__file__).parent.parent / "ZIYOU-Analytics-Documentacao.docx"

# ── Paleta de cores
AZUL_ESCURO  = RGBColor(0x07, 0x07, 0x10)   # fundo do dashboard
AZUL_MEDIO   = RGBColor(0x4F, 0x8E, 0xF7)   # accent azul
AMARELO      = RGBColor(0xFF, 0xE6, 0x00)   # accent amarelo
VERDE        = RGBColor(0x22, 0xD3, 0xA0)   # verde success
VERMELHO     = RGBColor(0xEF, 0x44, 0x44)   # erro
CINZA_CLARO  = RGBColor(0xF0, 0xF0, 0xFF)   # texto claro
CINZA_MEDIO  = RGBColor(0x88, 0x88, 0xAA)   # texto secundário
CINZA_FUNDO  = RGBColor(0xF5, 0xF5, 0xFA)   # fundo de código
PRETO        = RGBColor(0x1A, 0x1A, 0x2E)   # texto principal
BRANCO       = RGBColor(0xFF, 0xFF, 0xFF)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_border_left(para, color: str = "4F8EF7", size: int = 24):
    """Adiciona borda esquerda colorida ao parágrafo (estilo callout)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(size))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def add_run_with_style(para, text, bold=False, italic=False,
                       color=None, size=None, font_name=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:      run.font.color.rgb = color
    if size:       run.font.size = Pt(size)
    if font_name:  run.font.name = font_name
    return run


def heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if color:
            run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(22)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(13)
            run.font.bold = True
    return h


def body(doc, text, size=11, indent=0, color=None, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def numbered(doc, text, level=0, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def code_block(doc, text, title=None):
    """Bloco de código com fundo cinza e fonte monoespaçada."""
    if title:
        p = doc.add_paragraph()
        r = p.add_run(f"  {title}")
        r.font.size  = Pt(9)
        r.font.bold  = True
        r.font.color.rgb = AZUL_MEDIO
        p.paragraph_format.space_after = Pt(0)

    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(f"  {line}" if line else "")
        r.font.name  = 'Courier New'
        r.font.size  = Pt(9)
        r.font.color.rgb = PRETO
        # Fundo cinza simulado via shading no parágrafo
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F8')
        pPr.append(shd)

    # Linha em branco após
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def callout(doc, text, kind='info'):
    """Caixa de destaque com borda colorida."""
    colors = {'info': '4F8EF7', 'warn': 'FFE600', 'danger': 'EF4444', 'ok': '22D3A0'}
    labels = {'info': 'ℹ DICA', 'warn': '⚠ ATENÇÃO', 'danger': '✖ CUIDADO', 'ok': '✔ PRONTO'}
    col = colors.get(kind, '4F8EF7')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    set_para_border_left(p, col)

    rgb = {'info': AZUL_MEDIO, 'warn': RGBColor(0xFF,0xA5,0x00),
           'danger': VERMELHO, 'ok': VERDE}.get(kind, AZUL_MEDIO)

    r = p.add_run(labels.get(kind, 'NOTA') + '  ')
    r.font.bold  = True
    r.font.size  = Pt(10)
    r.font.color.rgb = rgb

    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = PRETO
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Cria tabela formatada com cabeçalho azul."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], '0F0F2E')
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(10)
                run.font.color.rgb = BRANCO

    # Linhas
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5FA'
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # Larguras
    if col_widths:
        for i, row in enumerate(t.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    doc.add_paragraph()
    return t


def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('─' * 80)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)
    r.font.size = Pt(8)


# ─────────────────────────────────────────────
# CAPA
# ─────────────────────────────────────────────

def build_cover(doc):
    # Espaço superior
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ZIYOU Analytics")
    r.font.size  = Pt(40)
    r.font.bold  = True
    r.font.color.rgb = PRETO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Dashboard de Mercado Livre")
    r2.font.size  = Pt(20)
    r2.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Documentação Completa")
    r3.font.size  = Pt(16)
    r3.font.bold  = True
    r3.font.color.rgb = PRETO

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Guia para quem nunca programou")
    r4.font.size  = Pt(13)
    r4.font.color.rgb = CINZA_MEDIO

    for _ in range(5):
        doc.add_paragraph()

    # Tabela de informações
    t = doc.add_table(rows=6, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    infos = [
        ("Dashboard online",  "https://brunnomark315-sys.github.io/ziyou-analytics/"),
        ("Repositório GitHub","https://github.com/brunnomark315-sys/ziyou-analytics"),
        ("Login de acesso",   "cliente"),
        ("Senha de acesso",   "Ziyou2026!"),
        ("Conta Mercado Livre","ZIYOUBRAND  (ID: 3277035412)"),
        ("Data do documento", "Maio 2026"),
    ]
    for i, (k, v) in enumerate(infos):
        row = t.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_bg(row.cells[0], '0F0F2E')
        set_cell_bg(row.cells[1], 'F5F5FA')
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(10)
                run.font.color.rgb = BRANCO
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    add_page_break(doc)


# ─────────────────────────────────────────────
# ÍNDICE (manual — Word atualiza automaticamente)
# ─────────────────────────────────────────────

def build_toc(doc):
    heading(doc, "Índice", level=1)

    sections = [
        ("1.", "O que é este projeto",                         "3"),
        ("2.", "Como tudo funciona — visão geral do fluxo",    "5"),
        ("3.", "Como criar conta no Mercado Livre Developers",  "6"),
        ("4.", "Como configurar a autenticação",                "8"),
        ("5.", "Como abrir o projeto no Claude Code",           "10"),
        ("6.", "Como funciona o MCP",                          "11"),
        ("7.", "Como os dados são atualizados automaticamente", "12"),
        ("8.", "Como publicar o dashboard online",              "14"),
        ("9.", "Como trocar login e senha",                     "16"),
        ("10.", "Como adicionar um novo cliente",               "17"),
        ("11.", "Principais erros e como resolver",             "19"),
        ("12.", "Se eu esquecer tudo — guia de retomada",       "21"),
    ]

    t = doc.add_table(rows=len(sections), cols=3)
    for i, (num, title, pg) in enumerate(sections):
        t.rows[i].cells[0].text = num
        t.rows[i].cells[1].text = title
        t.rows[i].cells[2].text = pg
        bg = 'FFFFFF' if i % 2 == 0 else 'F5F5FA'
        for j in range(3):
            set_cell_bg(t.rows[i].cells[j], bg)
            for para in t.rows[i].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
        for para in t.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = AZUL_MEDIO
        for para in t.rows[i].cells[2].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.color.rgb = CINZA_MEDIO

    callout(doc,
            "Este índice é uma referência visual. Para navegar no Word, use os "
            "Estilos de Título (Exibir → Painel de Navegação).",
            kind='info')
    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 1 — O QUE É ESTE PROJETO
# ─────────────────────────────────────────────

def section_what(doc):
    heading(doc, "1. O que é este projeto", level=1)

    body(doc, (
        "O ZIYOU Analytics Dashboard é uma página na internet que mostra, de forma "
        "visual e organizada, todos os dados importantes da loja ZIYOUBRAND no "
        "Mercado Livre: quantas visitas cada produto recebeu, quantas vendas foram "
        "feitas, qual é o estoque disponível e como está a saúde geral da loja."
    ))
    body(doc, (
        "Pense nele como um painel de controle de carro — mas em vez de mostrar "
        "velocidade e combustível, ele mostra visitas, vendas e estoque da sua loja."
    ))

    heading(doc, "1.1 O que o dashboard mostra", level=2)
    bullet(doc, "Health Score (Nota de Saúde): uma pontuação de 0 a 100 que diz se a loja está indo bem ou precisa de atenção.")
    bullet(doc, "KPIs (Números Principais): total de anúncios, visitas, vendas, receita, ticket médio e estoque.")
    bullet(doc, "Gráficos: visitas e vendas dia a dia nos últimos 30 dias.")
    bullet(doc, "Tabela de Anúncios: todos os produtos com preço, estoque, visitas e status.")
    bullet(doc, "Ranking: os produtos mais visitados.")
    bullet(doc, "Insights de IA: sugestões automáticas do que melhorar na loja.")
    bullet(doc, "Histórico: comparação com períodos anteriores.")

    heading(doc, "1.2 Como o dashboard se conecta ao Mercado Livre", level=2)
    body(doc, (
        "O Mercado Livre tem uma 'porta de serviço' chamada API (Application "
        "Programming Interface). É como um balcão especial onde um programa de "
        "computador pode pedir informações: 'quais são os meus produtos?' ou "
        "'quantas visitas tive hoje?'. O nosso sistema usa essa porta para buscar "
        "os dados da conta ZIYOUBRAND automaticamente."
    ))
    callout(doc,
            "Para usar essa porta, o Mercado Livre exige um 'crachá' eletrônico "
            "chamado token de acesso. Esse crachá é gerado uma vez e renovado "
            "automaticamente pelo sistema.",
            kind='info')

    heading(doc, "1.3 Como os dados chegam ao dashboard", level=2)
    body(doc, "O caminho dos dados é simples:")
    body(doc, (
        "  1. Todo dia às 07h da manhã (horário de Brasília: 04h), um robô "
        "automático acorda.\n"
        "  2. Ele se conecta ao Mercado Livre e busca todos os dados.\n"
        "  3. Salva esses dados em um arquivo chamado data.json.\n"
        "  4. Esse arquivo vai automaticamente para o GitHub (um serviço de "
        "armazenamento de código).\n"
        "  5. O GitHub publica o arquivo atualizado no site do dashboard.\n"
        "  6. Quando o cliente abre o site, ele vê os dados mais recentes."
    ), size=11)

    heading(doc, "1.4 Como o cliente acessa", level=2)
    body(doc, "O cliente acessa o dashboard em duas etapas:")
    numbered(doc, "Abre o link no navegador: https://brunnomark315-sys.github.io/ziyou-analytics/")
    numbered(doc, "Uma tela de login aparece pedindo a senha.")
    numbered(doc, "Digita a senha e clica em Entrar.")
    numbered(doc, "O dashboard abre mostrando todos os dados da loja.")

    make_table(doc,
        ["Informação", "Valor"],
        [
            ["Endereço do dashboard", "https://brunnomark315-sys.github.io/ziyou-analytics/"],
            ["Login (usuário)", "cliente"],
            ["Senha", "Ziyou2026!"],
            ["Dados atualizados", "Todo dia às 04:00 AM (Brasília)"],
            ["Conta Mercado Livre", "ZIYOUBRAND — ID 3277035412"],
        ],
        col_widths=[6, 11]
    )

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 2 — COMO TUDO FUNCIONA
# ─────────────────────────────────────────────

def section_flow(doc):
    heading(doc, "2. Como tudo funciona — visão geral do fluxo", level=1)

    body(doc, (
        "Para entender o projeto sem precisar ser programador, imagine que "
        "existem 5 'peças' que trabalham juntas:"
    ))

    code_block(doc, """\
┌─────────────────────────────────────────────────────────────┐
│                                                             │
│  MERCADO LIVRE  →  ROBÔ  →  GITHUB  →  SITE  →  CLIENTE   │
│                                                             │
│  (dados reais)    (busca)  (guarda)  (mostra)  (vê)        │
│                                                             │
└─────────────────────────────────────────────────────────────┘""")

    heading(doc, "Explicando cada peça:", level=2)

    body(doc, "MERCADO LIVRE", bold=True)
    body(doc, (
        "É onde estão os dados reais: produtos, visitas, vendas. "
        "O Mercado Livre guarda tudo isso e disponibiliza para quem tem permissão."
    ), indent=0.5)

    body(doc, "ROBÔ (GitHub Actions + fetch_data.py)", bold=True)
    body(doc, (
        "É um programa que roda automaticamente todo dia. Ele se conecta ao "
        "Mercado Livre, pega todos os dados e os organiza em um arquivo. "
        "Pense nele como um funcionário que trabalha de madrugada, sem precisar "
        "que ninguém acorde."
    ), indent=0.5)

    body(doc, "GITHUB", bold=True)
    body(doc, (
        "É um serviço gratuito da internet para guardar código e arquivos. "
        "É como um Google Drive específico para projetos de software. "
        "Nosso projeto fica guardado lá em: github.com/brunnomark315-sys/ziyou-analytics"
    ), indent=0.5)

    body(doc, "SITE (GitHub Pages)", bold=True)
    body(doc, (
        "O GitHub tem um serviço chamado GitHub Pages que transforma os arquivos "
        "do projeto em um site real acessível pela internet. É gratuito e automático — "
        "toda vez que o robô atualiza os dados, o site é atualizado junto."
    ), indent=0.5)

    body(doc, "CLIENTE", bold=True)
    body(doc, (
        "Qualquer pessoa com o link e a senha pode abrir o dashboard em qualquer "
        "computador ou celular, sem instalar nada."
    ), indent=0.5)

    callout(doc,
            "O único momento em que você precisa fazer algo manualmente é quando "
            "quiser mudar alguma configuração ou adicionar um novo cliente. "
            "No dia a dia, tudo roda sozinho.",
            kind='ok')

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 3 — MERCADO LIVRE DEVELOPERS
# ─────────────────────────────────────────────

def section_ml_dev(doc):
    heading(doc, "3. Como criar conta no Mercado Livre Developers", level=1)

    body(doc, (
        "Para que o nosso sistema consiga buscar dados do Mercado Livre "
        "automaticamente, você precisa criar um 'aplicativo' no site de "
        "desenvolvedores do ML. Esse aplicativo é como uma chave que dá permissão "
        "para o nosso sistema acessar a conta da loja."
    ))

    heading(doc, "3.1 Passo a passo para criar o aplicativo", level=2)

    numbered(doc, "Abra o navegador (Chrome, Firefox ou Safari).")
    numbered(doc, "Digite na barra de endereços: developers.mercadolivre.com.br")
    numbered(doc, "Clique em 'Entrar' no canto superior direito.")
    numbered(doc, "Use o mesmo e-mail e senha da conta ZIYOUBRAND no Mercado Livre.")
    numbered(doc, "Após entrar, clique em 'Criar aplicativo' ou 'Meus aplicativos'.")
    numbered(doc, "Clique no botão 'Criar novo aplicativo'.")

    body(doc, "Preencha o formulário com estas informações:", bold=True)
    make_table(doc,
        ["Campo", "O que digitar"],
        [
            ["Nome do aplicativo", "ZIYOU Analytics (ou qualquer nome que identifique)"],
            ["Descrição curta",    "Dashboard de analytics para a loja ZIYOU"],
            ["URL de redirecionamento", "https://localhost (exatamente assim)"],
            ["Escopos necessários",    "Marque: read_orders, read_listings, read_metrics"],
        ],
        col_widths=[5.5, 11.5]
    )

    callout(doc,
            "A URL de redirecionamento 'https://localhost' parece estranha mas é "
            "necessária para o processo de autenticação funcionar. Não é preciso "
            "ter um site em localhost — é apenas um endereço temporário que o ML "
            "usa para entregar o código de autorização.",
            kind='info')

    numbered(doc, "Clique em 'Salvar' ou 'Criar'.")
    numbered(doc, "O ML vai mostrar a tela do seu aplicativo criado.")

    heading(doc, "3.2 Como encontrar o App ID e a Chave Secreta", level=2)

    body(doc, (
        "Após criar o aplicativo, você verá dois códigos importantes na tela. "
        "Anote os dois em um lugar seguro:"
    ))

    make_table(doc,
        ["O que é", "Onde fica", "Exemplo (fictício)"],
        [
            ["App ID\n(também chamado Client ID)",
             "Logo após criar, na tela do app. Campo 'App ID' ou 'Client ID'.",
             "1234567890"],
            ["Chave Secreta\n(também chamada Client Secret)",
             "Na mesma tela, campo 'Secret Key' ou 'Client Secret'. Pode precisar clicar em 'Mostrar'.",
             "aB3xY9zK2mN7pQ1r"],
        ],
        col_widths=[4, 8, 5]
    )

    callout(doc,
            "CUIDADO: A Chave Secreta é como a senha do seu aplicativo. "
            "Nunca compartilhe com ninguém e não a coloque em lugares públicos "
            "como e-mails ou redes sociais.",
            kind='danger')

    heading(doc, "3.3 Como saber se deu certo", level=2)
    bullet(doc, "Você consegue ver o App ID (um número com 10+ dígitos).")
    bullet(doc, "Você consegue ver a Chave Secreta (uma combinação de letras e números).")
    bullet(doc, "O status do aplicativo aparece como 'Ativo'.")

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 4 — AUTENTICAÇÃO
# ─────────────────────────────────────────────

def section_auth(doc):
    heading(doc, "4. Como configurar a autenticação", level=1)

    heading(doc, "4.1 O que é OAuth e por que precisamos dele", level=2)
    body(doc, (
        "OAuth é um sistema de segurança que permite que um programa acesse sua "
        "conta em outro serviço sem que você precise digitar sua senha naquele "
        "programa. Em vez disso, você aprova o acesso diretamente no site do "
        "serviço (no caso, o Mercado Livre)."
    ))
    body(doc, (
        "É como quando você entra em algum site clicando em 'Entrar com Google'. "
        "Você nunca dá sua senha do Google para aquele site — você aprova o acesso "
        "diretamente no Google. O processo do OAuth para o Mercado Livre funciona "
        "da mesma forma."
    ))
    body(doc, "O resultado é um 'token de acesso' — uma chave eletrônica temporária que o nosso "
              "sistema usa para buscar dados sem precisar da sua senha a cada vez.")

    heading(doc, "4.2 Pré-requisitos antes de começar", level=2)
    bullet(doc, "Ter o App ID e a Chave Secreta (explicados na seção 3).")
    bullet(doc, "Ter o projeto na pasta: /Users/brunnomark/mercado-livre-analytics")
    bullet(doc, "Ter o Python instalado no Mac (verificar na seção 5).")

    heading(doc, "4.3 Passo a passo para autenticar", level=2)

    numbered(doc, "Abra o Terminal do Mac.")
    body(doc, "Para abrir o Terminal: pressione Command + Espaço, digite 'Terminal' e pressione Enter.", indent=1)

    numbered(doc, "Digite o comando abaixo e pressione Enter:")
    code_block(doc, "cd /Users/brunnomark/mercado-livre-analytics", title="No Terminal:")

    numbered(doc, "Digite o segundo comando para iniciar a autenticação:")
    code_block(doc, "python3 auth_setup.py", title="No Terminal:")

    numbered(doc, "O programa vai pedir o App ID. Digite o número e pressione Enter.")
    numbered(doc, "O programa vai pedir a Chave Secreta. Digite e pressione Enter.")
    numbered(doc, "Uma URL longa vai aparecer na tela. Copie essa URL.")
    numbered(doc, "Cole a URL no navegador e pressione Enter.")
    numbered(doc, "Uma tela do Mercado Livre vai aparecer pedindo para você fazer login e autorizar o aplicativo.")
    numbered(doc, "Após autorizar, o navegador vai tentar abrir uma página que não existe (localhost). Isso é normal!")
    numbered(doc, "Copie o endereço completo da barra do navegador (que começa com https://localhost?code=...).")
    numbered(doc, "Volte ao Terminal e cole esse endereço. Pressione Enter.")

    heading(doc, "4.4 O que deve aparecer quando der certo", level=2)
    code_block(doc, """\
✓ Token de acesso obtido com sucesso!
✓ Tokens salvos em: /Users/brunnomark/.mercado-livre-analytics/tokens.json
✓ Seller ID: 3277035412
✓ Conta: ZIYOUBRAND

Autenticação concluída. O sistema está pronto para buscar dados.""",
        title="Mensagem de sucesso esperada:")

    heading(doc, "4.5 Como saber se funcionou", level=2)
    bullet(doc, "A mensagem de sucesso apareceu no Terminal (como mostrado acima).")
    bullet(doc, "Existe um arquivo em: /Users/brunnomark/.mercado-livre-analytics/tokens.json")
    bullet(doc, "Ao executar o Claude Code, o MCP consegue buscar dados do ML.")

    callout(doc,
            "O token de acesso expira após 6 horas, mas o sistema renova "
            "automaticamente usando o refresh_token salvo no arquivo tokens.json. "
            "Você não precisa repetir esse processo frequentemente — só quando o "
            "refresh_token expirar (geralmente após 6 meses sem uso).",
            kind='info')

    heading(doc, "4.6 Onde as credenciais ficam guardadas", level=2)
    make_table(doc,
        ["Arquivo", "Onde fica", "O que contém"],
        [
            ["tokens.json",
             "/Users/brunnomark/.mercado-livre-analytics/tokens.json",
             "Token de acesso + refresh token + seller ID. NUNCA commitar no GitHub."],
            ["ML_CLIENT_ID",
             "GitHub Secrets (configurações do repositório)",
             "O App ID do Mercado Livre Developers."],
            ["ML_CLIENT_SECRET",
             "GitHub Secrets (configurações do repositório)",
             "A Chave Secreta do app ML Developers."],
            ["ML_TOKENS_JSON",
             "GitHub Secrets (configurações do repositório)",
             "Conteúdo do tokens.json para o robô automático usar."],
        ],
        col_widths=[4, 7, 6]
    )

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 5 — CLAUDE CODE
# ─────────────────────────────────────────────

def section_claude(doc):
    heading(doc, "5. Como abrir o projeto no Claude Code", level=1)

    body(doc, (
        "Claude Code é o programa que usamos para desenvolver e manter o projeto. "
        "Pense nele como um assistente inteligente que fica dentro do Terminal e "
        "sabe tudo sobre o projeto."
    ))

    heading(doc, "5.1 Pré-requisitos", level=2)
    bullet(doc, "Ter o Claude Code instalado (verifique digitando 'claude' no Terminal).")
    bullet(doc, "Estar na pasta correta do projeto.")
    bullet(doc, "Ter o arquivo .claude/settings.local.json configurado (já está no projeto).")

    heading(doc, "5.2 Como abrir o projeto", level=2)

    numbered(doc, "Abra o Terminal (Command + Espaço → 'Terminal' → Enter).")
    numbered(doc, "Entre na pasta do projeto:")
    code_block(doc, "cd /Users/brunnomark/mercado-livre-analytics")
    numbered(doc, "Inicie o Claude Code:")
    code_block(doc, "claude")
    numbered(doc, "Uma tela diferente do terminal vai aparecer — isso é o Claude Code.")

    heading(doc, "5.3 Como testar se está funcionando", level=2)
    body(doc, "No Claude Code, peça para ele verificar os dados:")
    code_block(doc, "Use get_reputation e me diga o nickname da conta", title="Digite no Claude Code:")
    body(doc, "O Claude deve responder com:")
    code_block(doc, """\
Nickname: ZIYOUBRAND
User ID:  3277035412
Status:   Ativo""", title="Resposta esperada:")

    callout(doc,
            "Se aparecer uma mensagem de erro sobre 'tokens não encontrados', "
            "significa que a autenticação precisa ser refeita (seção 4). "
            "Se aparecer 'MCP não conectado', verifique o arquivo "
            ".claude/settings.local.json.",
            kind='warn')

    heading(doc, "5.4 Comandos úteis no Claude Code", level=2)
    make_table(doc,
        ["O que você quer fazer", "O que digitar no Claude Code"],
        [
            ["Ver os produtos da loja",       "Use get_products e liste todos os anúncios ativos"],
            ["Ver as visitas do mês",          "Use get_visits para os últimos 30 dias"],
            ["Ver vendas recentes",            "Use get_orders e mostre os últimos 5 pedidos pagos"],
            ["Ver a reputação da loja",        "Use get_reputation e me mostre todos os dados"],
            ["Atualizar o dashboard",          "Execute scripts/fetch_data.py e atualize o site"],
            ["Gerar o PDF executivo",          "Execute scripts/generate_pdf.py"],
        ],
        col_widths=[7, 10]
    )

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 6 — MCP
# ─────────────────────────────────────────────

def section_mcp(doc):
    heading(doc, "6. Como funciona o MCP", level=1)

    heading(doc, "6.1 O que é o MCP", level=2)
    body(doc, (
        "MCP significa 'Model Context Protocol'. É uma forma de conectar o Claude "
        "(assistente de IA) a ferramentas e dados externos. Sem o MCP, o Claude "
        "só consegue conversar — com o MCP, ele consegue buscar dados reais do "
        "Mercado Livre, ler arquivos, executar código e muito mais."
    ))
    body(doc, (
        "Pense no MCP como uma 'tomada universal' que permite ao Claude se "
        "conectar a diferentes serviços como se fossem extensões."
    ))

    heading(doc, "6.2 O MCP deste projeto", level=2)
    body(doc, (
        "O arquivo server.py é o servidor MCP do projeto. Ele disponibiliza para "
        "o Claude as seguintes ferramentas:"
    ))
    make_table(doc,
        ["Ferramenta", "O que faz", "Exemplo de uso"],
        [
            ["get_products",       "Lista todos os anúncios da loja com preço e estoque",  "Use get_products e liste os ativos"],
            ["get_visits",         "Busca quantas visitas cada anúncio recebeu",            "Use get_visits dos últimos 30 dias"],
            ["get_orders",         "Mostra os pedidos feitos na loja",                      "Use get_orders com status paid"],
            ["get_sales",          "Resume o total de vendas e receita",                    "Use get_sales de abril a maio"],
            ["get_reputation",     "Mostra a reputação e métricas do vendedor",             "Use get_reputation"],
            ["get_financial_summary", "Mostra receita, comissões e lucro líquido",          "Use get_financial_summary"],
        ],
        col_widths=[4.5, 7, 5.5]
    )

    heading(doc, "6.3 Como verificar se o MCP está conectado", level=2)
    numbered(doc, "Abra o Claude Code na pasta do projeto.")
    numbered(doc, "Pergunte algo que usa o MCP:")
    code_block(doc, "Use get_reputation e me diga o nickname da conta")
    numbered(doc, "Se aparecer os dados da conta, o MCP está funcionando.")
    numbered(doc, "Se aparecer erro, veja os problemas comuns na seção 11.")

    heading(doc, "6.4 Onde o MCP é configurado", level=2)
    body(doc, "O MCP está configurado no arquivo:")
    code_block(doc, "/Users/brunnomark/mercado-livre-analytics/.claude/settings.local.json")
    body(doc, "Esse arquivo diz ao Claude Code para usar o server.py como servidor MCP quando estiver nessa pasta.")

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 7 — ATUALIZAÇÃO AUTOMÁTICA
# ─────────────────────────────────────────────

def section_auto(doc):
    heading(doc, "7. Como os dados são atualizados automaticamente", level=1)

    body(doc, (
        "O dashboard não precisa de ninguém para atualizar os dados. Um robô "
        "faz isso sozinho todo dia. Esse robô se chama GitHub Actions."
    ))

    heading(doc, "7.1 O que é GitHub Actions", level=2)
    body(doc, (
        "GitHub Actions é como um relógio despertador que acorda um robô em um "
        "horário específico para executar tarefas. O nosso robô acorda todo dia "
        "às 07h00 (horário UTC — equivale a 04h00 de Brasília) e faz o seguinte:"
    ))

    code_block(doc, """\
07:00 UTC — Robô acorda
   ↓
Conecta ao Mercado Livre com o token salvo
   ↓
Busca: produtos, visitas, pedidos, reputação, financeiro
   ↓
Calcula o Health Score (nota 0-100 da loja)
   ↓
Gera os Insights automáticos
   ↓
Salva tudo em public/data.json
   ↓
Gera o PDF executivo do relatório
   ↓
Faz upload para o GitHub
   ↓
GitHub Pages atualiza o site automaticamente
   ↓
Dashboard online mostra dados novos
   ↓
~07:10 UTC — Tudo concluído""")

    heading(doc, "7.2 Como verificar se o robô rodou", level=2)
    numbered(doc, "Abra o navegador e entre em: github.com/brunnomark315-sys/ziyou-analytics")
    numbered(doc, "Clique na aba 'Actions' (quarta opção no menu superior do repositório).")
    numbered(doc, "Você vai ver uma lista de execuções. A mais recente aparece no topo.")
    numbered(doc, "Se aparecer um círculo verde com ✓, o robô rodou com sucesso.")
    numbered(doc, "Se aparecer um círculo vermelho com ✗, houve um erro.")

    heading(doc, "7.3 O que fazer se o robô falhou", level=2)
    body(doc, "Clique na execução com erro para ver o que aconteceu:")
    bullet(doc, "Erro de token: a autenticação expirou. Refaça o auth_setup.py (seção 4) e atualize o secret ML_TOKENS_JSON.")
    bullet(doc, "Erro de API: o Mercado Livre está fora do ar. Aguarde e o robô tenta novamente amanhã.")
    bullet(doc, "Erro de permissão: os GitHub Secrets podem ter sido apagados. Reconfigure (seção 7.4).")

    heading(doc, "7.4 Como configurar os segredos no GitHub", level=2)
    body(doc, (
        "Os segredos são informações confidenciais (como senhas e tokens) que o "
        "robô precisa para funcionar. Eles ficam guardados no GitHub de forma segura."
    ))
    numbered(doc, "Entre em: github.com/brunnomark315-sys/ziyou-analytics")
    numbered(doc, "Clique em 'Settings' (Configurações) — último item do menu superior.")
    numbered(doc, "No menu lateral esquerdo, clique em 'Secrets and variables' → 'Actions'.")
    numbered(doc, "Clique em 'New repository secret'.")
    numbered(doc, "Crie os três segredos abaixo:")

    make_table(doc,
        ["Nome do Segredo", "O que colocar como valor"],
        [
            ["ML_CLIENT_ID",      "O App ID do Mercado Livre Developers (número de 10+ dígitos)"],
            ["ML_CLIENT_SECRET",  "A Chave Secreta do app ML Developers"],
            ["ML_TOKENS_JSON",    "Todo o conteúdo do arquivo tokens.json (abra o arquivo, selecione tudo e cole)"],
        ],
        col_widths=[5, 12]
    )

    body(doc, "Para ver o conteúdo do tokens.json, abra o Terminal e digite:")
    code_block(doc, "cat ~/.mercado-livre-analytics/tokens.json")
    body(doc, "Copie tudo que aparecer (começando com { e terminando com }) e cole no valor do segredo ML_TOKENS_JSON.")

    heading(doc, "7.5 Como atualizar manualmente", level=2)
    body(doc, "Para atualizar os dados sem esperar o horário automático:")
    numbered(doc, "Entre em: github.com/brunnomark315-sys/ziyou-analytics/actions")
    numbered(doc, "Clique no workflow 'Update Dashboard Data'.")
    numbered(doc, "Clique no botão 'Run workflow' (lado direito).")
    numbered(doc, "Clique no botão verde 'Run workflow' que aparece.")
    numbered(doc, "Aguarde cerca de 2-3 minutos e o dashboard será atualizado.")

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 8 — PUBLICAÇÃO
# ─────────────────────────────────────────────

def section_deploy(doc):
    heading(doc, "8. Como publicar o dashboard online", level=1)

    heading(doc, "8.1 O que é o GitHub Pages", level=2)
    body(doc, (
        "GitHub Pages é um serviço gratuito do GitHub que transforma arquivos "
        "de um repositório em um site real na internet. É como um 'hospedeiro de "
        "site gratuito'. O nosso dashboard já está publicado lá."
    ))

    heading(doc, "8.2 O repositório atual", level=2)
    make_table(doc,
        ["Informação", "Valor"],
        [
            ["Repositório",       "github.com/brunnomark315-sys/ziyou-analytics"],
            ["Branch publicada",  "main"],
            ["Pasta publicada",   "public/ (contém index.html, app.js, style.css, data.json)"],
            ["URL do site",       "https://brunnomark315-sys.github.io/ziyou-analytics/"],
            ["HTTPS ativo",       "Sim (certificado automático)"],
        ],
        col_widths=[5, 12]
    )

    heading(doc, "8.3 Como publicar do zero (se precisar recriar)", level=2)

    body(doc, "Passo 1: criar o repositório no GitHub", bold=True)
    numbered(doc, "Entre em github.com e faça login.")
    numbered(doc, "Clique no '+' no canto superior direito → 'New repository'.")
    numbered(doc, "Nome: ziyou-analytics")
    numbered(doc, "Marque 'Public' (público).")
    numbered(doc, "Clique em 'Create repository'.")

    body(doc, "Passo 2: enviar os arquivos", bold=True)
    body(doc, "No Terminal, dentro da pasta do projeto:")
    code_block(doc, """\
git init
git add public/ scripts/ clients/ .github/ auth.py auth_setup.py server.py pyproject.toml .gitignore .env.example wrangler.toml
git commit -m "primeiro commit"
git branch -M main
git remote add origin https://github.com/brunnomark315-sys/ziyou-analytics.git
git push -u origin main""")

    body(doc, "Passo 3: ativar o GitHub Pages", bold=True)
    numbered(doc, "No repositório, clique em 'Settings'.")
    numbered(doc, "No menu lateral, clique em 'Pages'.")
    numbered(doc, "Em 'Source', selecione 'GitHub Actions'.")
    numbered(doc, "O workflow pages.yml já está no repositório e fará o deploy automaticamente.")

    heading(doc, "8.4 Como configurar um domínio próprio (analytics.ziyou.com.br)", level=2)
    body(doc, (
        "Para usar um endereço personalizado em vez do github.io, é preciso "
        "fazer dois ajustes: um no provedor de domínio (onde o domínio foi comprado "
        "ou configurado) e um no GitHub."
    ))

    body(doc, "Parte 1: adicionar o registro DNS", bold=True)
    body(doc, (
        "Entre no painel de controle do seu domínio (Cloudflare, Registro.br, etc.) "
        "e adicione este registro:"
    ))
    make_table(doc,
        ["Tipo", "Nome", "Valor", "TTL"],
        [
            ["CNAME", "analytics", "brunnomark315-sys.github.io", "Auto"],
        ],
        col_widths=[2.5, 3.5, 8.5, 2.5]
    )
    callout(doc,
            "Se estiver usando Cloudflare: mantenha o Proxy desativado (ícone de "
            "nuvem cinza, não laranja). Caso contrário, o GitHub não consegue "
            "verificar o domínio.",
            kind='warn')

    body(doc, "Parte 2: criar arquivo CNAME no projeto", bold=True)
    body(doc, "Crie um arquivo chamado CNAME (sem extensão) dentro da pasta public/ com o conteúdo:")
    code_block(doc, "analytics.ziyou.com.br", title="Conteúdo do arquivo public/CNAME:")

    body(doc, "Depois, faça commit e push:")
    code_block(doc, """\
git add public/CNAME
git commit -m "adicionar dominio customizado"
git push""")

    body(doc, "Parte 3: configurar no GitHub", bold=True)
    numbered(doc, "No repositório, vá em Settings → Pages.")
    numbered(doc, "Em 'Custom domain', digite: analytics.ziyou.com.br")
    numbered(doc, "Clique em 'Save'.")
    numbered(doc, "Aguarde alguns minutos para o certificado HTTPS ser emitido.")

    callout(doc,
            "Enquanto o DNS não propagar (pode levar até 48h), o site ficará "
            "disponível apenas na URL do GitHub. Após propagar, ambas as URLs "
            "funcionam, mas o GitHub redireciona para o domínio customizado.",
            kind='info')

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 9 — TROCAR LOGIN E SENHA
# ─────────────────────────────────────────────

def section_password(doc):
    heading(doc, "9. Como trocar login e senha", level=1)

    body(doc, (
        "O login do dashboard usa um sistema seguro onde a senha nunca é guardada "
        "como texto simples — apenas uma 'impressão digital' matemática dela (chamada hash). "
        "Para trocar a senha, você gera uma nova impressão digital e atualiza o arquivo de configuração."
    ))

    heading(doc, "9.1 Passo a passo para trocar a senha", level=2)

    numbered(doc, "Abra o Terminal.")
    numbered(doc, "Gere a 'impressão digital' da nova senha (substitua minhasenha pela senha real):")
    code_block(doc, 'python3 -c "import hashlib; print(hashlib.sha256(b\'minhasenha\').hexdigest())"',
               title="No Terminal:")
    body(doc, "Vai aparecer um código longo como:", indent=0.5)
    code_block(doc, "a94a8fe5ccb19ba61c4c0873d391e987982fbbd3",
               title="Exemplo de hash gerado:")

    numbered(doc, "Copie esse código longo (selecione e pressione Command+C).")
    numbered(doc, "Abra o arquivo de configuração:")
    code_block(doc, "open /Users/brunnomark/mercado-livre-analytics/public/config.json",
               title="No Terminal:")
    numbered(doc, "Encontre a linha que começa com 'password_hash' e substitua o código pelo novo.")
    body(doc, "Antes:", bold=True, indent=0.5)
    code_block(doc, '"password_hash": "d191117e129cc1cb56b2ecf71228f7442c06de809013833e54dd92819a46d48d"')
    body(doc, "Depois (com o novo hash):", bold=True, indent=0.5)
    code_block(doc, '"password_hash": "SEU_NOVO_HASH_AQUI"')
    numbered(doc, "Salve o arquivo (Command+S).")
    numbered(doc, "Envie para o GitHub para o site ser atualizado:")
    code_block(doc, """\
cd /Users/brunnomark/mercado-livre-analytics
git add public/config.json
git commit -m "atualizar senha do dashboard"
git push""")

    heading(doc, "9.2 Senha atual de acesso", level=2)
    make_table(doc,
        ["Campo", "Valor atual"],
        [
            ["Login",   "cliente"],
            ["Senha",   "Ziyou2026!"],
            ["Hash SHA-256", "d191117e129cc1cb56b2ecf71228f7442c06de809013833e54dd92819a46d48d"],
        ],
        col_widths=[4, 13]
    )

    heading(doc, "9.3 Como desativar o login (acesso livre)", level=2)
    body(doc, "Para deixar o dashboard sem senha (não recomendado para dados reais), edite config.json e mude:")
    code_block(doc, '"password_protection": false')
    callout(doc,
            "Desativar o login significa que qualquer pessoa com o link pode "
            "ver todos os dados da loja. Recomendamos sempre manter o login ativo.",
            kind='danger')

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 10 — NOVO CLIENTE
# ─────────────────────────────────────────────

def section_new_client(doc):
    heading(doc, "10. Como adicionar um novo cliente", level=1)

    body(doc, (
        "O projeto foi construído para funcionar com múltiplos clientes. "
        "Para cada cliente novo (outra loja do Mercado Livre), você cria uma "
        "pasta separada dentro de 'clients/' com as configurações específicas. "
        "Cada cliente tem seu próprio dashboard, senha e dados."
    ))

    heading(doc, "10.1 Estrutura de pastas para múltiplos clientes", level=2)
    code_block(doc, """\
mercado-livre-analytics/
└── clients/
    ├── ziyou/           ← cliente atual
    │   ├── config.json  ← configurações da ZIYOU
    │   └── snapshots/   ← histórico de dados da ZIYOU
    │
    └── novocliente/     ← novo cliente (criar)
        ├── config.json  ← configurações do novo cliente
        └── snapshots/   ← histórico do novo cliente""")

    heading(doc, "10.2 Passo a passo para adicionar um novo cliente", level=2)

    body(doc, "Etapa 1: criar a pasta do cliente", bold=True)
    code_block(doc, """\
cd /Users/brunnomark/mercado-livre-analytics
mkdir -p clients/nomecliente/snapshots
touch clients/nomecliente/snapshots/.gitkeep""",
        title="No Terminal (substitua 'nomecliente' pelo ID do cliente):")

    body(doc, "Etapa 2: criar o arquivo de configuração", bold=True)
    body(doc, "Copie o config.json da ZIYOU como modelo:")
    code_block(doc, "cp clients/ziyou/config.json clients/nomecliente/config.json")
    body(doc, "Abra o novo config.json e altere:")
    make_table(doc,
        ["Campo", "O que mudar"],
        [
            ["client_id",        "Trocar 'ziyou' para o ID do novo cliente (ex: 'lojaxyz')"],
            ["client_name",      "Nome que aparece no dashboard (ex: 'Loja XYZ')"],
            ["logo_text",        "Letra inicial do nome (ex: 'L')"],
            ["logo_gradient_from / to", "Cores em hexadecimal do logo (ex: '#0066FF', '#0044CC')"],
            ["password_hash",    "Hash SHA-256 da senha do novo cliente (seção 9.1)"],
            ["category_map",     "Mapa de categorias específico da conta ML do cliente"],
        ],
        col_widths=[5, 12]
    )

    body(doc, "Etapa 3: autenticar com a conta ML do novo cliente", bold=True)
    code_block(doc, "python3 auth_setup.py", title="No Terminal:")
    body(doc, "Use as credenciais (App ID e Chave Secreta) do aplicativo ML do novo cliente.")

    body(doc, "Etapa 4: buscar os primeiros dados", bold=True)
    code_block(doc, "python3 scripts/fetch_data.py --client nomecliente",
               title="No Terminal:")

    body(doc, "Etapa 5: publicar o dashboard do cliente", bold=True)
    body(doc, "Para publicar em um repositório separado para cada cliente:")
    numbered(doc, "Crie um novo repositório GitHub com o nome do cliente.")
    numbered(doc, "Copie os arquivos da pasta public/ para o novo repositório.")
    numbered(doc, "Configure os GitHub Secrets com as credenciais ML do novo cliente.")
    numbered(doc, "Ative o GitHub Pages no novo repositório.")

    callout(doc,
            "No futuro, o sistema pode ser expandido para gerar dashboards separados "
            "para cada cliente automaticamente, usando o GitHub Actions com 'matrix' "
            "(execução em paralelo para múltiplos clientes). Por enquanto, crie um "
            "repositório separado para cada cliente.",
            kind='info')

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 11 — ERROS
# ─────────────────────────────────────────────

def section_errors(doc):
    heading(doc, "11. Principais erros e como resolver", level=1)

    body(doc, "Esta seção reúne os problemas reais que encontramos durante o desenvolvimento do projeto.")

    make_table(doc,
        ["Problema", "Causa provável", "Como resolver"],
        [
            [
                "Dashboard abre sem pedir login",
                "1) Usuário já tinha feito login na mesma sessão do navegador (sessionStorage guardado). "
                "2) config.json não carregou corretamente.",
                "Abra o dashboard em uma aba anônima (Command+Shift+N no Chrome). "
                "Isso limpa o sessionStorage. Se continuar sem login, verifique se config.json tem "
                "'password_protection': true e se o hash está correto."
            ],
            [
                "Site exibe '301 Moved Permanently' e não abre",
                "O domínio customizado (analytics.ziyou.com.br) está configurado no GitHub Pages "
                "mas o registro DNS (CNAME) ainda não foi adicionado.",
                "Adicione o registro CNAME no provedor do domínio (seção 8.4) OU remova o domínio "
                "customizado temporariamente via GitHub Pages Settings → Pages → remover o Custom Domain."
            ],
            [
                "Erro: 'Tokens não encontrados. Execute auth_setup.py'",
                "O arquivo tokens.json não existe ou expirou.",
                "Execute python3 auth_setup.py no Terminal e complete o processo de autenticação (seção 4)."
            ],
            [
                "Erro: 'ML_CLIENT_ID e ML_CLIENT_SECRET são obrigatórios'",
                "As variáveis de ambiente não estão definidas.",
                "Para rodar localmente: export ML_CLIENT_ID=seu_id e export ML_CLIENT_SECRET=sua_chave. "
                "Para o robô automático: verifique os GitHub Secrets (seção 7.4)."
            ],
            [
                "GitHub Actions falha com 'The certificate does not exist yet'",
                "O GitHub tentou configurar o domínio customizado antes do DNS propagar.",
                "Aguarde a propagação do DNS (até 48h). Depois acesse Settings → Pages e clique em Save novamente."
            ],
            [
                "get_orders retorna lista vazia",
                "A conta não tem pedidos no status solicitado (ex: 'paid') no período consultado.",
                "Verifique se existem pedidos na conta do ML. Se a conta for nova, pode não ter pedidos ainda."
            ],
            [
                "Health Score está muito baixo (abaixo de 30)",
                "A conta tem poucos anúncios com visitas, sem vendas ou estoque parado.",
                "Siga as ações prioritárias exibidas no próprio dashboard. O score melhora conforme as métricas melhoram."
            ],
            [
                "gh: command not found (no Terminal)",
                "O GitHub CLI não está instalado.",
                "Execute: brew install gh — depois: gh auth login"
            ],
            [
                "python3: command not found",
                "Python não está instalado ou não está no PATH.",
                "No Mac, execute: brew install python3 — ou instale pelo site python.org"
            ],
            [
                "git push: rejeitado (permission denied)",
                "O token do GitHub expirou ou a autenticação mudou.",
                "Execute: gh auth login — e configure as credenciais novamente."
            ],
            [
                "Dashboard mostra dados desatualizados",
                "O robô automático não rodou hoje ou falhou.",
                "Verifique em github.com/brunnomark315-sys/ziyou-analytics/actions se o último workflow deu sucesso. "
                "Se falhou, clique em 'Run workflow' para rodar manualmente (seção 7.5)."
            ],
            [
                "config.json retorna erro 404",
                "O arquivo não está na pasta public/ ou não foi feito push.",
                "Verifique se o arquivo existe: ls public/config.json — se não existir, recrie e faça push."
            ],
        ],
        col_widths=[5, 5.5, 6.5]
    )

    add_page_break(doc)


# ─────────────────────────────────────────────
# SEÇÃO 12 — SE EU ESQUECER TUDO
# ─────────────────────────────────────────────

def section_resume(doc):
    heading(doc, "12. Se eu esquecer tudo — guia de retomada", level=1)

    body(doc, (
        "Esta seção é para quando você voltar ao projeto depois de meses sem tocá-lo "
        "e precisar retomar tudo do zero. Siga a lista em ordem."
    ))

    heading(doc, "12.1 Checklist de retomada (em ordem)", level=2)

    body(doc, "VERIFICAR 1: O site ainda está no ar?", bold=True)
    body(doc, "Abra o navegador e entre em:", indent=0.5)
    code_block(doc, "https://brunnomark315-sys.github.io/ziyou-analytics/")
    body(doc, "Se aparecer a tela de login → SITE OK, vá para Verificar 2.", indent=0.5)
    body(doc, "Se aparecer erro 404 → o GitHub Pages está desativado. Vá em Settings → Pages e ative.", indent=0.5)
    doc.add_paragraph()

    body(doc, "VERIFICAR 2: O robô automático está funcionando?", bold=True)
    body(doc, "Abra:", indent=0.5)
    code_block(doc, "https://github.com/brunnomark315-sys/ziyou-analytics/actions")
    body(doc, "Veja a data do último workflow 'Update Dashboard Data'.", indent=0.5)
    body(doc, "Se o último foi há mais de 2 dias e tem ícone vermelho → robô falhando. Veja Verificar 3.", indent=0.5)
    doc.add_paragraph()

    body(doc, "VERIFICAR 3: Os tokens ML ainda são válidos?", bold=True)
    body(doc, "Abra o Terminal e entre na pasta:", indent=0.5)
    code_block(doc, "cd /Users/brunnomark/mercado-livre-analytics")
    body(doc, "Verifique se o arquivo de tokens existe:", indent=0.5)
    code_block(doc, "cat ~/.mercado-livre-analytics/tokens.json")
    body(doc, "Se mostrar um JSON com access_token → tokens existem. Mas podem estar expirados.", indent=0.5)
    body(doc, "Se mostrar 'No such file' → precisa refazer a autenticação (seção 4).", indent=0.5)
    doc.add_paragraph()

    body(doc, "VERIFICAR 4: Os GitHub Secrets estão configurados?", bold=True)
    body(doc, "Vá em:", indent=0.5)
    code_block(doc, "github.com/brunnomark315-sys/ziyou-analytics → Settings → Secrets → Actions")
    body(doc, "Devem existir: ML_CLIENT_ID, ML_CLIENT_SECRET, ML_TOKENS_JSON", indent=0.5)
    body(doc, "Se algum estiver faltando → reconfigure (seção 7.4).", indent=0.5)
    doc.add_paragraph()

    body(doc, "VERIFICAR 5: O Claude Code ainda funciona?", bold=True)
    code_block(doc, """\
cd /Users/brunnomark/mercado-livre-analytics
claude""", title="No Terminal:")
    body(doc, "Depois, no Claude Code, teste:", indent=0.5)
    code_block(doc, "Use get_reputation e me diga o nickname")
    body(doc, "Se retornar ZIYOUBRAND → tudo funcionando.", indent=0.5)
    doc.add_paragraph()

    heading(doc, "12.2 Se precisar reautenticar do zero", level=2)
    body(doc, "Use quando os tokens expirarem completamente:")
    numbered(doc, "Tenha em mãos: App ID e Chave Secreta do ML Developers.")
    numbered(doc, "Execute: python3 auth_setup.py")
    numbered(doc, "Complete o processo de autorização no navegador.")
    numbered(doc, "Copie o novo tokens.json e atualize o GitHub Secret ML_TOKENS_JSON.")
    numbered(doc, "Acione o robô manualmente para testar (seção 7.5).")

    heading(doc, "12.3 Se precisar reinstalar tudo em um Mac novo", level=2)
    make_table(doc,
        ["O que instalar", "Como instalar", "Verificar com"],
        [
            ["Homebrew (gerenciador de pacotes)", "Acesse brew.sh e siga as instruções", "brew --version"],
            ["Python 3", "brew install python3", "python3 --version"],
            ["Git", "brew install git", "git --version"],
            ["GitHub CLI", "brew install gh", "gh --version"],
            ["Claude Code", "Acesse claude.ai/code e instale", "claude --version"],
        ],
        col_widths=[5, 7, 5]
    )

    body(doc, "Depois de instalar, clone o projeto do GitHub:")
    code_block(doc, """\
gh auth login
gh repo clone brunnomark315-sys/ziyou-analytics
cd ziyou-analytics""")

    heading(doc, "12.4 Informações essenciais para não perder", level=2)
    callout(doc,
            "Guarde estas informações em um lugar seguro (cofre de senhas, papel guardado, etc.):",
            kind='warn')
    make_table(doc,
        ["Informação", "Valor / Onde encontrar"],
        [
            ["URL do dashboard",          "https://brunnomark315-sys.github.io/ziyou-analytics/"],
            ["Repositório GitHub",         "github.com/brunnomark315-sys/ziyou-analytics"],
            ["Login GitHub",               "brunnomark315-sys"],
            ["Pasta local do projeto",     "/Users/brunnomark/mercado-livre-analytics"],
            ["App ID do ML",               "Ver no ML Developers ou no GitHub Secret ML_CLIENT_ID"],
            ["Chave Secreta do ML",        "Ver no ML Developers ou no GitHub Secret ML_CLIENT_SECRET"],
            ["Tokens ML",                  "~/.mercado-livre-analytics/tokens.json (renovar se expirar)"],
            ["Login do dashboard",         "cliente"],
            ["Senha do dashboard",         "Ziyou2026!"],
            ["ID do vendedor ML",          "3277035412 (ZIYOUBRAND)"],
        ],
        col_widths=[6, 11]
    )

    heading(doc, "12.5 Fluxo de desenvolvimento — como continuar", level=2)
    body(doc, "Para fazer qualquer alteração no projeto:")
    numbered(doc, "Abra o Terminal e entre na pasta: cd /Users/brunnomark/mercado-livre-analytics")
    numbered(doc, "Inicie o Claude Code: claude")
    numbered(doc, "Descreva o que quer fazer em linguagem natural (não precisa saber programar).")
    numbered(doc, "Revise o que o Claude propõe antes de confirmar.")
    numbered(doc, "Após as alterações, publique: git add . → git commit -m 'descrição' → git push")
    numbered(doc, "O site é atualizado automaticamente em 1-2 minutos.")

    callout(doc,
            "Você não precisa entender o código para manter o projeto. "
            "O Claude Code explica cada passo e executa as mudanças por você. "
            "Descreva o que quer em português simples e ele faz o resto.",
            kind='ok')

    add_page_break(doc)


# ─────────────────────────────────────────────
# RODAPÉ DA ÚLTIMA PÁGINA
# ─────────────────────────────────────────────

def build_final(doc):
    heading(doc, "Glossário rápido", level=1)
    make_table(doc,
        ["Termo", "Significado simples"],
        [
            ["API",            "Porta de comunicação entre sistemas. O ML usa para entregar dados."],
            ["Token",          "Chave eletrônica temporária que dá acesso a uma conta."],
            ["Hash",           "Impressão digital matemática de um texto (ex: de uma senha)."],
            ["GitHub",         "Serviço de armazenamento de código na internet."],
            ["GitHub Pages",   "Serviço do GitHub que transforma código em site."],
            ["GitHub Actions", "Robô automático do GitHub que executa tarefas em horários programados."],
            ["MCP",            "Protocolo que conecta o Claude (IA) a ferramentas externas."],
            ["OAuth",          "Sistema de autorização segura entre serviços (sem compartilhar senha)."],
            ["CNAME",          "Tipo de registro DNS que aponta um domínio para outro endereço."],
            ["DNS",            "Sistema de endereços da internet (transforma nomes em IPs)."],
            ["JSON",           "Formato de arquivo para guardar dados estruturados (como uma planilha em texto)."],
            ["sessionStorage", "Memória temporária do navegador que some quando a aba é fechada."],
            ["Health Score",   "Nota de 0 a 100 que mede a saúde do marketplace da loja."],
            ["KPI",            "Key Performance Indicator — métricas principais de desempenho."],
            ["Snapshot",       "Fotografia dos dados em um momento específico para comparação futura."],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Documentação gerada automaticamente pelo projeto ZIYOU Analytics · Maio 2026\n"
        "github.com/brunnomark315-sys/ziyou-analytics"
    )
    r.font.size  = Pt(9)
    r.font.color.rgb = CINZA_MEDIO
    r.italic = True


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    doc = Document()

    # Configurar margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Configurar estilo padrão
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = PRETO
    style.paragraph_format.space_after = Pt(6)

    # Construir documento
    build_cover(doc)
    build_toc(doc)
    section_what(doc)
    section_flow(doc)
    section_ml_dev(doc)
    section_auth(doc)
    section_claude(doc)
    section_mcp(doc)
    section_auto(doc)
    section_deploy(doc)
    section_password(doc)
    section_new_client(doc)
    section_errors(doc)
    section_resume(doc)
    build_final(doc)

    doc.save(str(OUT))
    print(f"✓ Documentação gerada → {OUT}")
    print(f"  Tamanho: {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
